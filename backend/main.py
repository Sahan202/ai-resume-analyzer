from io import BytesIO
import json
import os

from docx import Document
from dotenv import load_dotenv
from fastapi import FastAPI, File, Form, HTTPException, UploadFile
from fastapi.middleware.cors import CORSMiddleware
from openai import OpenAI, OpenAIError
from pypdf import PdfReader

load_dotenv()

MAX_FILE_SIZE = 10 * 1024 * 1024
ALLOWED_EXTENSIONS = {".pdf", ".docx"}
SKILLS = [
    "python", "javascript", "typescript", "react", "node.js", "fastapi", "django",
    "sql", "mongodb", "postgresql", "aws", "docker", "git", "github", "figma",
    "excel", "power bi", "tableau", "machine learning", "data analysis", "html", "css",
]

app = FastAPI(title="ResumeAI API", version="0.1.0")

app.add_middleware(
    CORSMiddleware,
    allow_origins=["http://localhost:5173"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)


def extract_resume_text(content: bytes, extension: str) -> str:
    """Extract readable text from a PDF or DOCX resume file."""
    try:
        if extension == ".pdf":
            reader = PdfReader(BytesIO(content))
            return "\n".join(page.extract_text() or "" for page in reader.pages).strip()

        document = Document(BytesIO(content))
        paragraphs = [paragraph.text for paragraph in document.paragraphs]
        table_cells = [cell.text for table in document.tables for row in table.rows for cell in row.cells]
        return "\n".join(paragraphs + table_cells).strip()
    except Exception as error:
        raise HTTPException(status_code=400, detail="We could not read this resume file.") from error


def analyze_resume(resume_text: str, job_description: str) -> dict:
    """A transparent local scoring baseline; replaceable with an AI provider later."""
    resume = resume_text.lower()
    job = job_description.lower()
    required_skills = [skill for skill in SKILLS if skill in job]
    matching_skills = [skill for skill in required_skills if skill in resume]
    missing_skills = [skill for skill in required_skills if skill not in resume]
    skill_score = round((len(matching_skills) / len(required_skills)) * 55) if required_skills else 28

    experience_terms = ("experience", "worked", "developer", "engineer", "intern", "manager")
    has_experience = any(term in resume for term in experience_terms)
    needs_experience = any(term in job for term in experience_terms)
    experience_score = 25 if has_experience else 8
    education_terms = ("bachelor", "degree", "university", "college", "diploma", "education")
    education_score = 15 if any(term in resume for term in education_terms) else 5
    formatting_score = 5 if len(resume_text) > 600 else 1
    ats_score = min(100, skill_score + experience_score + education_score + formatting_score)

    suggestions = []
    if missing_skills:
        suggestions.append("Add relevant evidence for these missing job skills: " + ", ".join(missing_skills[:4]) + ".")
    if not has_experience and needs_experience:
        suggestions.append("Add a work, internship, freelance, or project experience section with measurable outcomes.")
    if education_score < 15:
        suggestions.append("Include your education, degree or relevant certifications in a clearly labeled section.")
    if len(resume_text) < 600:
        suggestions.append("Add more detail to projects and achievements, using action verbs and numbers where possible.")
    if not suggestions:
        suggestions.append("Strong match. Tailor your summary and project bullet points to use the exact language in the job description.")

    return {
        "ats_score": ats_score,
        "matching_skills": matching_skills,
        "missing_skills": missing_skills,
        "experience_match": "Strong" if has_experience and needs_experience else "Needs more evidence",
        "education_match": "Present" if education_score == 15 else "Not clearly found",
        "suggestions": suggestions,
    }


def configured_ai_provider() -> str | None:
    if os.getenv("GROQ_API_KEY"):
        return "groq"
    if os.getenv("USE_OPENAI", "false").lower() == "true" and os.getenv("OPENAI_API_KEY"):
        return "openai"
    return None


def analyze_with_ai(resume_text: str, job_description: str, provider: str | None) -> dict | None:
    """Return a tailored assessment from a configured server-side AI provider."""
    if not provider:
        return None

    prompt = f"""You are an expert resume reviewer. Compare the resume and job description below.
Return ONLY valid JSON with exactly these fields:
ats_score (integer 0-100), matching_skills (array of strings), missing_skills (array of strings),
experience_match (short string), education_match (short string), suggestions (array of 3-5 specific strings).
Do not invent qualifications. Assess only evidence present in the resume. Be fair and concise.

RESUME:
{resume_text[:18000]}

JOB DESCRIPTION:
{job_description[:12000]}"""

    try:
        if provider == "groq":
            client = OpenAI(
                api_key=os.environ["GROQ_API_KEY"],
                base_url="https://api.groq.com/openai/v1",
            )
            model = os.getenv("GROQ_MODEL", "llama-3.3-70b-versatile")
        else:
            client = OpenAI(api_key=os.environ["OPENAI_API_KEY"])
            model = os.getenv("OPENAI_MODEL", "gpt-4o-mini")
        if provider == "groq":
            response = client.chat.completions.create(
                model=model,
                messages=[
                    {"role": "system", "content": "You return valid JSON only."},
                    {"role": "user", "content": prompt},
                ],
                response_format={"type": "json_object"},
            )
            output_text = response.choices[0].message.content or ""
        else:
            response = client.responses.create(model=model, input=prompt)
            output_text = response.output_text
        result = json.loads(output_text)
        required_fields = {
            "ats_score", "matching_skills", "missing_skills", "experience_match", "education_match", "suggestions"
        }
        if not required_fields.issubset(result):
            raise ValueError("AI response did not contain the expected analysis fields.")
        result["ats_score"] = max(0, min(100, int(result["ats_score"])))
        return result
    except (OpenAIError, json.JSONDecodeError, TypeError, ValueError) as error:
        raise HTTPException(status_code=502, detail=f"AI analysis failed: {error}") from error


@app.get("/health")
async def health_check():
    return {"status": "ok"}


@app.post("/api/resumes/analyze")
async def upload_resume(
    resume: UploadFile = File(...), job_description: str = Form(...)
):
    """Accepts the upload. Text extraction and AI scoring are added in later stages."""
    filename = resume.filename or "resume"
    extension = "." + filename.rsplit(".", 1)[-1].lower() if "." in filename else ""

    if extension not in ALLOWED_EXTENSIONS:
        raise HTTPException(status_code=400, detail="Only PDF and DOCX files are supported.")
    if not job_description.strip():
        raise HTTPException(status_code=400, detail="A job description is required.")

    content = await resume.read()
    if len(content) > MAX_FILE_SIZE:
        raise HTTPException(status_code=400, detail="Resume must be 10 MB or smaller.")

    resume_text = extract_resume_text(content, extension)
    if not resume_text:
        raise HTTPException(
            status_code=400,
            detail="No readable text was found. Please upload a text-based PDF or DOCX file.",
        )

    ai_provider = configured_ai_provider()
    ai_analysis = analyze_with_ai(resume_text, job_description, ai_provider)
    analysis = ai_analysis or analyze_resume(resume_text, job_description)
    return {
        "message": "Resume text extracted successfully.",
        "filename": filename,
        "file_size": len(content),
        "job_description_length": len(job_description.strip()),
        "resume_text": resume_text,
        "resume_text_length": len(resume_text),
        "analysis": analysis,
        "analysis_source": ai_provider if ai_analysis else "rule_based",
    }
